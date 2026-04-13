"""Gera DOCX com metodologia da interface web para dissertacao."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(12)
    return p


def main():
    out = Path(
        r"C:\Users\obser\OneDrive\Documentos\Gabriel\ESTAGIO\PESQUISA\modelo_web_hidrologia"
    ) / "Metodologia_Interface_Web_Modelo_Hidrologico.docx"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    title = doc.add_heading(
        "Metodologia de implementacao da interface web do modelo hidrologico urbano",
        0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Este texto descreve, de forma sequencial, o processo de concepcao, "
        "desenvolvimento, validacao e disponibilizacao em nuvem da interface web "
        "associada ao modelo de quantificacao do escoamento superficial e conversao "
        "da vazao em nivel d'agua no canal, de modo a subsidiar a secao metodologica "
        "de uma dissertacao de mestrado."
    )

    doc.add_heading("1. Justificativa e objetivos da interface web", level=1)
    doc.add_paragraph(
        "A modelagem hidrologica urbana baseada em balanco de massa e na equacao de "
        "Manning, quando implementada exclusivamente em planilhas eletronicas, apresenta "
        "limitacoes em termos de reproducibilidade, escalabilidade e experiencia do "
        "usuario. A criacao de uma interface web tem como objetivos: (i) reunir entradas, "
        "parametros e saidas em um unico ambiente interativo; (ii) reduzir erros de "
        "manuseio de formulas e referencias cruzadas entre abas; (iii) permitir "
        "visualizacao imediata de hidrogramas e metricas de desempenho; e (iv) facilitar "
        "a divulgacao e o uso do modelo por terceiros sem instalacao de software "
        "proprietario alem de um navegador."
    )

    doc.add_heading("2. Escolha tecnologica e arquitetura geral", level=1)
    doc.add_paragraph(
        "Optou-se pela linguagem Python pela maturidade do ecossistema cientifico "
        "(bibliotecas NumPy e Pandas para arranjos e tabelas) e pela existencia de "
        "frameworks de interface rapida. O framework Streamlit foi adotado por permitir "
        "a construcao de aplicacoes web locais e implantaveis em nuvem com poucas linhas "
        "de codigo, integrando formularios, tabelas editaveis e graficos interativos."
    )
    doc.add_paragraph(
        "A arquitetura do aplicativo segue o modelo de pagina unica com navegacao lateral "
        "(menu por radio): cada secao corresponde a um conjunto funcional (simulacao, "
        "validacao, sugestao de calibracao, instrucoes e glossario), mantendo o nucleo "
        "numerico em funcoes puras separadas da camada de apresentacao."
    )
    doc.add_paragraph(
        "Para graficos de series temporais utilizou-se a biblioteca Plotly Express / "
        "Graph Objects, por oferecer escalas legiveis, interatividade (zoom, pan) e "
        "boa integracao com Streamlit."
    )

    doc.add_heading("2.1 Ambiente de desenvolvimento assistido por inteligencia artificial (Cursor)", level=2)
    doc.add_paragraph(
        "A implementacao do codigo-fonte, a estruturacao dos modulos da interface, "
        "refatoracoes, geracao de documentacao auxiliar (por exemplo, README e scripts "
        "de apoio) e a elaboracao iterativa deste relato metodologico foram realizadas "
        "com apoio do ambiente integrado de desenvolvimento Cursor, aplicativo que integra "
        "recursos de inteligencia artificial generativa ao fluxo de edicao de codigo e texto."
    )
    doc.add_paragraph(
        "O Cursor permite, entre outras funcoes, sugestao e revisao de trechos de "
        "programa, explicacao de erros, busca contextual no projeto e assistencia na "
        "redacao tecnica, sempre sob supervisao do pesquisador, que define requisitos, "
        "valida resultados numericos frente a metodologia hidrologica e assume a "
        "responsabilidade cientifica pelo modelo e pela interface. Registra-se explicitamente "
        "esse uso para transparencia metodologica, em linha com boas praticas de ciencia "
        "aberta e de declaracao de ferramentas computacionais contemporaneas em trabalhos "
        "academicos."
    )

    doc.add_heading("3. Implementacao numerica no ambiente web", level=1)
    doc.add_paragraph(
        "O motor de calculo replica a formulacao metodologica do modelo: discretizacao "
        "temporal em passos regulares; para cada area de contribuicao (AC), calculo da "
        "chuva efetiva a partir da precipitacao e da taxa de infiltracao equivalente; "
        "balanco da lamina superficial; propagacao com coeficiente de Manning "
        "equivalente na superficie; soma das contribuicoes de montante quando ha mais "
        "de uma AC a jusante; selecao da vazao no exutorio; conversao da vazao em "
        "profundidade na secao retangular do canal por iteracao da equacao de Manning "
        "no canal."
    )
    doc.add_paragraph(
        "As unidades de uso do solo (US) podem ser informadas com areas e parametros "
        "por classe; o sistema calcula medias ponderadas por area para infiltracao e "
        "rugosidade equivalente por AC, alinhando-se ao procedimento de ponderacao "
        "descrito na metodologia documental do estudo."
    )

    doc.add_heading("4. Modulos da interface e fluxo de uso", level=1)
    items = [
        "Simulacao: parametros gerais (passo de tempo, duracao do evento), geometria e "
        "conectividade das ACs, tabela de US, serie de precipitacao em mm/h, parametros "
        "do canal; execucao da simulacao e exibicao de hidrogramas de vazao e nivel, "
        "tabelas de saida e exportacao para planilha Excel.",
        "Validacao: importacao ou digitacao de niveis observados no tempo alinhados aos "
        "instantes simulados; grafico observado versus simulado; indicadores Nash-Sutcliffe "
        "(NSE), coeficiente de determinacao baseado na correlacao de Pearson (R2), RMSE, "
        "MAE e vies medio.",
        "Sugestao de calibracao: tabelas de referencia com valores iniciais e faixas de "
        "variacao de infiltracao e Manning por classe de uso do solo, e coeficientes de "
        "Manning para tipos de revestimento de canal, servindo de guia ao usuario na "
        "parametrizacao.",
        "Instrucoes e descricao das variaveis: texto orientativo e glossario dos simbolos "
        "utilizados na interface.",
    ]
    for i, t in enumerate(items, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(t).font.size = Pt(12)

    doc.add_heading("5. Execucao local (ambiente de desenvolvimento)", level=1)
    doc.add_paragraph(
        "O prototipo e testado em ambiente local mediante interpretador Python e "
        "instalacao das dependencias listadas em arquivo requirements.txt (Streamlit, "
        "Pandas, NumPy, Plotly, openpyxl). O comando streamlit run app.py inicia um "
        "servidor HTTP na maquina, acessivel em navegador no endereco local (por exemplo, "
        "http://localhost:8501). Essa etapa e utilizada para depuracao, ajuste de layout "
        "e validacao dos resultados antes da publicacao."
    )

    doc.add_heading("6. Controle de versao e repositorio remoto", level=1)
    doc.add_paragraph(
        "O codigo-fonte foi versionado com Git. O repositorio remoto no GitHub concentra "
        "os arquivos necessarios ao deploy (app.py, requirements.txt, README, configuracao "
        "opcional em .streamlit), excluindo-se artefatos desnecessarios ao runtime "
        "mediante arquivo .gitignore (ambientes virtuais, caches Python, segredos locais)."
    )
    doc.add_paragraph(
        "Atualizacoes do aplicativo seguem o fluxo: alteracao local, commit descritivo, "
        "push para o branch principal; o servico de hospedagem reconstrói a aplicacao a "
        "partir do ultimo commit aceito."
    )

    doc.add_heading("7. Disponibilizacao publica (computacao em nuvem)", level=1)
    doc.add_paragraph(
        "Para disponibilizar o modelo sem exigir terminal ou instalacao na maquina do "
        "usuario final, utilizou-se o Streamlit Community Cloud (ou servico equivalente). "
        "O fluxo consiste em: (1) criar repositorio publico no GitHub com o codigo da "
        "aplicacao; (2) associar a conta GitHub ao Streamlit Cloud; (3) configurar o "
        "deploy indicando repositorio, branch (tipicamente main), e caminho do arquivo "
        "principal app.py na raiz do repositorio; (4) aguardar o build e obter URL publica "
        "no dominio streamlit.app."
    )
    doc.add_paragraph(
        "Ressalta-se que aplicacoes publicas em plano gratuito estao sujeitas a limites "
        "de uso e que dados inseridos na interface transitam pelo servidor do provedor; "
        "para dados sensiveis recomenda-se politicas institucionais ou hospedagem privada."
    )

    doc.add_heading("8. Limitacoes e extensoes futuras", level=1)
    doc.add_paragraph(
        "A interface atual prioriza clareza e um unico evento de chuva por simulacao; "
        "extensoes possiveis incluem autenticacao de usuarios, persistencia de cenarios em "
        "banco de dados, importacao automatizada de series pluviometricas de APIs e "
        "parametrizacao assistida por otimizacao formal com base na aba de validacao."
    )

    doc.add_heading("9. Referencias complementares (documentacao tecnica)", level=1)
    refs = [
        "Documentacao Streamlit: https://docs.streamlit.io",
        "Documentacao Plotly Python: https://plotly.com/python/",
        "Streamlit Community Cloud: https://streamlit.io/cloud",
    ]
    for r in refs:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(r).font.size = Pt(12)

    doc.add_paragraph()
    p_end = doc.add_paragraph(
        "Texto gerado para incorporacao e edicao conforme normas da instituição e orientacao."
    )
    p_end.italic = True

    doc.save(out)
    print(f"Arquivo criado: {out}")


if __name__ == "__main__":
    main()
